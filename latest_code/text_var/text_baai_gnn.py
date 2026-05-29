"""
GNN 增强版法规文档语义拓扑分析器。

针对 **法律法规 / 规章制度** 类文档设计：
- CCAR 121-R8、民航规章、法律文书等
- 文档结构层级分明（编→章→节→条）
- 条款之间存在大量 **交叉引用**（如"依据第XX条"）

核心思路：
1. 将文档构建为 **异构图**：段落为节点，引用/层级/邻接为边
2. 使用 **GCN/GAT** 在图上传播语义，生成 GNN 增强嵌入
3. GNN 输出作为 PHATE 降维 + 聚类 + 分段的输入

优势：
- 引用关系让跨章节的关联条款在语义空间中自然靠近
- 层级结构保持文档组织的先验知识
- 相比 EMA 时序平滑，更适配非叙事型法规文本
"""

from __future__ import annotations

import argparse
import importlib
import json
import logging
import os
import re
import sys
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Set, Tuple

import numpy as np
import plotly.graph_objects as go
import torch
import torch.nn.functional as F

from plotly.subplots import make_subplots
from sentence_transformers import SentenceTransformer
from sklearn.neighbors import NearestNeighbors

# ---------------------------------------------------------------------------
# Logging 配置
# ---------------------------------------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)-7s | %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# scikit-learn 兼容性补丁
# ---------------------------------------------------------------------------
try:
    import sklearn.utils.validation as _skval

    if getattr(_skval.check_array, "_is_patched", False):
        importlib.reload(_skval)
    if not hasattr(_skval, "_original_check_array"):
        _skval._original_check_array = _skval.check_array

    def _patched_check_array(
        array: Any, *args: Any, **kwargs: Any
    ) -> Any:
        """移除 force_all_finite 参数以兼容新版 numpy."""
        kwargs.pop("force_all_finite", None)
        return _skval._original_check_array(array, *args, **kwargs)

    _patched_check_array._is_patched = True  # type: ignore[attr-defined]
    if _skval.check_array is not _patched_check_array:
        _skval.check_array = _patched_check_array
        logger.info("已应用 scikit-learn 兼容性补丁")
except ImportError:
    pass

# ---------------------------------------------------------------------------
# 依赖检测
# ---------------------------------------------------------------------------

_TDA_AVAILABLE: bool
try:
    import gtda  # noqa: F401
    from gtda.homology import VietorisRipsPersistence  # noqa: F401

    _TDA_AVAILABLE = True
except ImportError:
    _TDA_AVAILABLE = False
    logger.warning("未检测到 gtda，TDA 模块将跳过")

_PHATE_AVAILABLE: bool
try:
    import phate  # noqa: F401

    _PHATE_AVAILABLE = True
except ImportError:
    _PHATE_AVAILABLE = False
    logger.warning("未检测到 phate，PHATE 模块将跳过")

_HDBSCAN_AVAILABLE: bool
try:
    import hdbscan  # noqa: F401

    _HDBSCAN_AVAILABLE = True
except ImportError:
    _HDBSCAN_AVAILABLE = False
    logger.warning("未检测到 hdbscan，离群分析将跳过")

_LEIDEN_AVAILABLE: bool
try:
    import igraph as ig  # noqa: F401
    import leidenalg  # noqa: F401

    _LEIDEN_AVAILABLE = True
except ImportError:
    _LEIDEN_AVAILABLE = False
    logger.warning("未检测到 igraph/leidenalg，将回退 HDBSCAN 聚类")

_TORCH_GEOMETRIC_AVAILABLE: bool
try:
    import torch_geometric  # noqa: F401

    _TORCH_GEOMETRIC_AVAILABLE = True
except ImportError:
    _TORCH_GEOMETRIC_AVAILABLE = False
    logger.warning("未检测到 torch_geometric，GNN 模块将禁用")


# ---------------------------------------------------------------------------
# 自定义异常
# ---------------------------------------------------------------------------


class DocumentError(Exception):
    """文档读取或解析相关的错误基类。"""


class DocumentNotFoundError(DocumentError):
    """输入文档不存在。"""


class UnsupportedFormatError(DocumentError):
    """不支持的文档格式。"""


class DocumentParseError(DocumentError):
    """文档结构解析失败。"""


class GNNTrainError(RuntimeError):
    """GNN 训练失败。"""


# ---------------------------------------------------------------------------
# 全局配置
# ---------------------------------------------------------------------------


@dataclass
class Config:
    """分析器的全部可调参数，集中管理便于实验追踪。

    GNN 模式下，根据文档的引用和层级结构构建图，
    用 GCN/GAT 传播语义，替代 EMA 时序平滑。

    Attributes:
        input_file_path: 输入文档的绝对路径。
        doc_title: 文档标题，用于缓存目录名和报告标题。
        sbert_model_name: Sentence-BERT 模型名称，默认 ``BAAI/bge-m3``。
        device: 计算设备，自动检测 ``mps`` / ``cuda`` / ``cpu``。
        gnn_hidden_dim: GNN 隐藏层维度。
        gnn_num_layers: GNN 层数（2~3 通常足够）。
        gnn_dropout: Dropout 比率。
        gnn_model: GNN 模型类型，``'GCN'`` 或 ``'GAT'``。
        gnn_lr: 学习率。
        gnn_epochs: 最大训练轮数。
        gnn_early_stop_patience: 早停耐心轮数（``None`` 禁用）。
        gnn_val_ratio: 验证集比例（0.0 ~ 0.3）。
        graph_use_citation: 是否使用引用边（如"第XX条"）。
        graph_use_adjacency: 是否使用相邻段落边。
        graph_adj_k: 相邻边滑动窗口大小。
        phate_n_components: PHATE 降维目标维度。
        phate_knn: PHATE k-NN 邻居数。
        phate_knn_dist: PHATE 距离度量。
        phate_n_pca: PHATE PCA 预降维维度（``None`` 自动）。
        phate_mds: PHATE MDS 类型。
        phate_mds_solver: PHATE MDS 求解器。
        enable_tda: 是否启用 TDA 拓扑分析。
    """

    # --- 输入 ---
    input_file_path: str = "/Users/terrysun/Desktop/CCAR 121-R8.docx"
    doc_title: str = "document"

    # --- 编码 ---
    sbert_model_name: str = "BAAI/bge-m3"
    device: str = field(
        default_factory=lambda: (
            "mps"
            if torch.backends.mps.is_available()
            else "cuda" if torch.cuda.is_available() else "cpu"
        )
    )

    # --- GNN 参数 ---
    gnn_hidden_dim: int = 256
    gnn_num_layers: int = 2
    gnn_dropout: float = 0.2
    gnn_model: str = "GAT"
    gnn_lr: float = 0.01
    gnn_epochs: int = 200
    gnn_early_stop_patience: Optional[int] = 20
    gnn_val_ratio: float = 0.1

    # --- 图构建参数 ---
    graph_use_citation: bool = True
    graph_use_adjacency: bool = True
    graph_adj_k: int = 3

    # --- PHATE ---
    phate_n_components: int = 3
    phate_knn: int = 7
    phate_knn_dist: str = "cosine"
    phate_n_pca: Optional[int] = None
    phate_mds: str = "metric"
    phate_mds_solver: str = "smacof"

    # --- 离群 / 图聚类 ---
    outlier_k_min: int = 6
    outlier_k_max: int = 18
    leiden_resolution: float = 1.0
    hdbscan_min_cluster_size: int = 5
    hdbscan_min_samples: int = 1

    # --- 开关 ---
    enable_tda: bool = False

    def __post_init__(self) -> None:
        """校验参数合法性。"""
        if self.gnn_num_layers < 1:
            raise ValueError("gnn_num_layers 必须 ≥ 1")
        if not 0.0 <= self.gnn_dropout <= 1.0:
            raise ValueError("gnn_dropout 须在 [0, 1] 范围内")
        if self.gnn_model not in ("GCN", "GAT"):
            raise ValueError(f"不支持的 GNN 模型: {self.gnn_model}，仅支持 GCN / GAT")
        if not 0.0 <= self.gnn_val_ratio <= 0.3:
            raise ValueError("gnn_val_ratio 须在 [0, 0.3] 范围内")


# ==================== 工具类 ====================


class SimpleTimer:
    """上下文管理器，用于计时并输出耗时信息。

    Usage::

        with SimpleTimer("数据处理"):
            data = process(...)
    """

    def __init__(self, name: str) -> None:
        """初始化计时器。

        Args:
            name: 计时任务的名称，用于日志输出。
        """
        self.name: str = name
        self.start: float = 0.0

    def __enter__(self) -> SimpleTimer:
        """进入上下文，记录开始时间。"""
        self.start = time.perf_counter()
        logger.info("▶️ [开始] %s ...", self.name)
        return self

    def __exit__(self, *args: Any) -> None:
        """离开上下文，输出耗时信息。"""
        elapsed = time.perf_counter() - self.start
        logger.info("✅ [完成] %s | 耗时: %.2fs", self.name, elapsed)


# ==================== 模块 1：文件读取 ====================


def get_script_dir() -> str:
    """返回本脚本所在目录的绝对路径。

    Returns:
        脚本所在目录的绝对路径字符串。
    """
    try:
        return os.path.dirname(os.path.abspath(__file__))
    except NameError:
        return os.getcwd()


def extract_text_from_docx(path: str) -> str:
    """使用 python-docx 从 .docx 文件中提取文本。

    Args:
        path: .docx 文件的绝对路径。

    Returns:
        合并后的文本字符串（段落以换行符分隔）。

    Raises:
        DocumentParseError: 读取或解析失败。
    """
    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(path)
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as exc:
        raise DocumentParseError(f"读取 .docx 失败: {exc}") from exc


def extract_text_from_doc(path: str) -> str:
    """在 macOS 上使用 ``textutil`` 从 .doc 文件中提取文本。

    依赖于 macOS 系统自带的 ``textutil`` 命令行工具。

    Args:
        path: .doc 文件的绝对路径。

    Returns:
        提取出的文本字符串。

    Raises:
        DocumentParseError: 转换或读取失败。
    """
    import subprocess
    import tempfile

    try:
        with tempfile.NamedTemporaryFile(
            suffix=".txt", delete=False, mode="w"
        ) as tmp:
            tmp_path: str = tmp.name
        subprocess.run(
            ["textutil", "-convert", "txt", "-output", tmp_path, path],
            check=True,
            capture_output=True,
        )
        with open(tmp_path, "r", encoding="utf-8") as f:
            text: str = f.read()
        os.unlink(tmp_path)
        return text
    except Exception as exc:
        # 确保清理临时文件
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)
        raise DocumentParseError(
            f"读取 .doc 失败（需 macOS textutil）: {exc}"
        ) from exc


def read_text_file(path: str) -> str:
    """根据文件扩展名读取文本内容，支持 .txt / .docx / .doc。

    Args:
        path: 文件绝对路径。

    Returns:
        提取出的文本字符串。

    Raises:
        UnsupportedFormatError: 不受支持的文件格式。
        DocumentNotFoundError: 文件不存在。
        DocumentParseError: 文件读取/解析失败。
    """
    filepath = Path(path)
    if not filepath.exists():
        raise DocumentNotFoundError(f"文件不存在: {path}")

    ext: str = filepath.suffix.lower()

    if ext == ".txt":
        try:
            return filepath.read_text(encoding="utf-8")
        except Exception as exc:
            raise DocumentParseError(f"读取 .txt 失败: {exc}") from exc

    if ext == ".docx":
        return extract_text_from_docx(path)

    if ext == ".doc":
        return extract_text_from_doc(path)

    raise UnsupportedFormatError(
        f"不支持的文件格式: {ext}（仅支持 .txt / .docx / .doc）"
    )


# ==================== 模块 2：文字分段 ====================


def split_text_into_segments(text: str) -> List[Dict[str, str]]:
    """将输入文字按句子边界分割为 segment 列表。

    支持中英文混合文本：
    - 中文：按 。！？分割，并尝试保留引号内的完整性
    - 英文：按 ``. ! ?`` + 空格 + 大写字母 分割
    - 连续换行也作为分割边界
    - 若分割后不足 3 段，则回退到按换行分割

    Args:
        text: 原始文本字符串。

    Returns:
        每个元素为 ``{"text": <句子>}`` 的列表。
    """
    t: str = text.strip()

    t = re.sub(r'([。！？])(?!["」』》\）\)】\s]*[」』》\）\)】])', r"\1\n", t)
    t = re.sub(r"([.!?])\s+(?=[A-Z\"])", r"\1\n", t)
    t = re.sub(r"\n\s*\n", "\n", t)

    sentences: List[str] = [s.strip() for s in t.split("\n") if s.strip()]
    segments: List[Dict[str, str]] = [{"text": s} for s in sentences]

    if len(segments) < 3:
        lines: List[str] = [
            l.strip() for l in text.strip().split("\n") if l.strip()
        ]
        segments = [{"text": line} for line in lines]

    return segments


# ==================== 模块 3：图构建（核心 GNN 部分）====================


def extract_article_refs(text: str) -> List[str]:
    """从文本中提取引用的条款编号（如 121.47, 121.161 等）。

    Args:
        text: 单段文本。

    Returns:
        引用的条款编号列表（如 ``['121.47', '121.161']``）。
    """
    pattern: str = r"第\s*([\d]+(?:\.\d+)?)\s*条"
    return re.findall(pattern, text)


def parse_document_structure(
    texts: List[str], cfg: Config
) -> Tuple[List[Tuple[int, int, str]], Dict[int, str]]:
    """解析文档结构和引用关系。

    功能：
    - 识别层级标签：第X编、第X章、第X节、第X条
    - 构建引用边：段落中的"第XXX条" → 目标段落
    - 构建邻接边：相邻段落滑动窗口

    Args:
        texts: 文档段落列表。
        cfg: 配置对象，决定边构建策略。

    Returns:
        (边列表, 段落元数据字典)
        每条边为 *(源节点索引, 目标节点索引, 边类型)*，
        边类型为 ``'citation'`` / ``'adjacent'`` 之一。
    """
    n: int = len(texts)
    edges: List[Tuple[int, int, str]] = []
    meta: Dict[int, str] = {}

    # --- 识别层级标签 ---
    level_pattern: re.Pattern = re.compile(
        r"^第[一二三四五六七八九十百千\d]+[编章节条]"
    )
    for i, t in enumerate(texts):
        m = level_pattern.match(t.strip())
        if m:
            meta[i] = m.group()

    # --- 引用边 ---
    if cfg.graph_use_citation:
        article_idx: Dict[str, int] = {}
        for i, t in enumerate(texts):
            for ref in extract_article_refs(t):
                if t.strip().startswith(f"第{ref}条"):
                    article_idx[ref] = i

        for i, t in enumerate(texts):
            for ref in extract_article_refs(t):
                target: Optional[int] = article_idx.get(ref)
                if target is not None and target != i:
                    edges.append((i, target, "citation"))

    # --- 邻接边 ---
    if cfg.graph_use_adjacency:
        k: int = cfg.graph_adj_k
        for i in range(n):
            left: int = max(0, i - k)
            right: int = min(n, i + k + 1)
            for j in range(left, right):
                if i != j:
                    edges.append((i, j, "adjacent"))

    return edges, meta


def build_pyg_graph(
    emb: np.ndarray,
    edges: List[Tuple[int, int, str]],
    device: str,
) -> "torch_geometric.data.Data":
    """构建 PyTorch Geometric 图对象。

    Args:
        emb: 形状 ``(n, d)`` 的原始 embedding。
        edges: 边列表 *(源, 目标, 类型)*。
        device: 目标设备（``'cpu'`` / ``'mps'`` / ``'cuda'``）。

    Returns:
        PyG 的 ``Data`` 对象，包含 ``x``, ``edge_index``, ``edge_attr``。

    Raises:
        ImportError: torch_geometric 未安装。
    """
    from torch_geometric.data import Data

    n: int = emb.shape[0]
    edge_set: Set[Tuple[int, int]] = set()
    edge_attrs: List[int] = []
    edge_type_map: Dict[str, int] = {
        "citation": 0,
        "adjacent": 1,
        "hierarchy": 2,
    }

    for s, t, etype in edges:
        if (s, t) not in edge_set:
            edge_set.add((s, t))
            edge_attrs.append(edge_type_map.get(etype, 0))

    if not edge_set:
        # 保底：至少连接相邻节点
        edge_set = {(i, i + 1) for i in range(n - 1)}
        edge_set.update({(i + 1, i) for i in range(n - 1)})
        edge_attrs = [1] * len(edge_set)

    edge_index: torch.Tensor = (
        torch.tensor(list(edge_set), dtype=torch.long, device=device)
        .t()
        .contiguous()
    )
    edge_attr: torch.Tensor = torch.tensor(
        edge_attrs, dtype=torch.long, device=device
    )
    x: torch.Tensor = torch.tensor(emb, dtype=torch.float, device=device)

    return Data(x=x, edge_index=edge_index, edge_attr=edge_attr)


# ==================== 模块 4：GNN 模型 ====================


class GNNEncoder(torch.nn.Module):
    """图神经网络编码器，用于在文档图上传播语义。

    支持 GCN 和 GAT 两种模型。
    输入为 SBERT embedding，输出为 GNN 增强后的 embedding（同维度）。
    所有隐藏层均使用 **残差连接** 和 LayerNorm 以保证训练稳定性。

    Attributes:
        model_type: 模型类型 ``'GCN'`` 或 ``'GAT'``。
        num_layers: 总层数。
    """

    def __init__(
        self,
        in_dim: int,
        hidden_dim: int,
        out_dim: int,
        num_layers: int = 2,
        dropout: float = 0.2,
        model_type: str = "GCN",
    ) -> None:
        """初始化 GNN 编码器。

        Args:
            in_dim: 输入特征维度。
            hidden_dim: 隐藏层维度。
            out_dim: 输出特征维度（推荐与 ``in_dim`` 一致）。
            num_layers: 总层数（≥ 1）。
            dropout: Dropout 比率。
            model_type: ``'GCN'`` 或 ``'GAT'``。

        Raises:
            ValueError: 不支持的模型类型。
        """
        super().__init__()
        self.dropout = dropout
        self.model_type = model_type
        self.num_layers = num_layers

        if model_type == "GAT":
            from torch_geometric.nn import GATConv
            conv_cls = GATConv
        elif model_type == "GCN":
            from torch_geometric.nn import GCNConv
            conv_cls = GCNConv
        else:
            raise ValueError(f"不支持的模型类型: {model_type}")

        self.convs: torch.nn.ModuleList = torch.nn.ModuleList()
        self.norms: torch.nn.ModuleList = torch.nn.ModuleList()
        self.skip_proj: Optional[torch.nn.Module] = None

        # 如果维度不一致，需要投影层做残差连接
        if in_dim != hidden_dim:
            self.skip_proj = torch.nn.Linear(in_dim, hidden_dim)

        # 第一层: in_dim → hidden_dim
        self.convs.append(conv_cls(in_dim, hidden_dim))
        self.norms.append(torch.nn.LayerNorm(hidden_dim))

        # 中间层: hidden_dim → hidden_dim
        for _ in range(num_layers - 2):
            self.convs.append(conv_cls(hidden_dim, hidden_dim))
            self.norms.append(torch.nn.LayerNorm(hidden_dim))

        # 输出层
        if num_layers >= 2:
            self.convs.append(conv_cls(hidden_dim, out_dim))
        else:
            self.convs = torch.nn.ModuleList([conv_cls(in_dim, out_dim)])

    def forward(
        self, x: torch.Tensor, edge_index: torch.Tensor
    ) -> torch.Tensor:
        """前向传播。

        Args:
            x: 形状 ``(n, in_dim)`` 的节点特征。
            edge_index: 形状 ``(2, m)`` 的边索引。

        Returns:
            形状 ``(n, out_dim)`` 的 GNN 增强特征。
        """
        for i, conv in enumerate(self.convs):
            identity: torch.Tensor = x
            x = conv(x, edge_index)

            if i < len(self.convs) - 1:
                # 残差连接（仅当维度匹配时）
                if self.skip_proj is not None and i == 0:
                    identity = self.skip_proj(identity)
                if x.shape == identity.shape:
                    x = x + identity
                x = self.norms[i](x)
                x = F.relu(x)
                x = F.dropout(x, p=self.dropout, training=self.training)
        return x


# ==================== 模块 5：核心分析器 ====================


class LogicAnalyzer:
    """GNN 增强核心分析器。

   完整流程：
   ``SBERT 编码 → 图构建 → GNN 传播 → PHATE → TDA → 聚类``

    Attributes:

        texts: 原始文本段落列表。
        embeddings_gpu: GPU 上的 GNN 增强 embedding 张量。
        embeddings_cpu: CPU numpy 副本。
        results: 各步骤输出结果的字典。
        gnn_model: 训练好的 GNNEncoder 实例（如果已执行过 GNN 传播）。
    """

    def __init__(
        self,
        embeddings_tensor: torch.Tensor,
        texts: List[str],
    ) -> None:
        """初始化分析器。

        Args:
            embeddings_tensor: 形状 ``(n, d)`` 的 float32 张量。
            texts: 原始文本段落列表。
        """
        self.texts: List[str] = texts
        self.embeddings_original: np.ndarray = (
            embeddings_tensor.detach().cpu().numpy()
        )
        self.embeddings_gpu: torch.Tensor = embeddings_tensor
        self.embeddings_cpu: np.ndarray = (
            embeddings_tensor.cpu().numpy()
            + np.random.normal(0, 1e-5, embeddings_tensor.shape)
        )
        self.results: Dict[str, Any] = {}
        self.gnn_model: Optional[GNNEncoder] = None

    # ------------------------------------------------------------------
    # 图构建 + GNN 传播（替代 EMA 平滑）
    # ------------------------------------------------------------------

    def run_gnn_propagation(self, cfg: Config) -> None:
        """构建文档图，运行 GNN 生成增强 embedding。

        流程：
        1. 解析文档结构和引用关系
        2. 构建 PyG ``Data`` 对象
        3. 划分训练/验证集，训练 GNN（自监督重建）
        4. 用训练好的 GNN 输出 + 残差连接作为增强 embedding

        Args:
            cfg: 配置对象。

        Raises:
            GNNTrainError: GNN 训练失败。
        """
        if not _TORCH_GEOMETRIC_AVAILABLE:
            logger.warning("torch_geometric 不可用，跳过 GNN！")
            return

        device: torch.device = torch.device(cfg.device)
        n, d = self.embeddings_gpu.shape

        with SimpleTimer("解析文档结构"):
            edges, meta = parse_document_structure(self.texts, cfg)
            n_citation: int = sum(
                1 for _, _, t in edges if t == "citation"
            )
            n_adjacent: int = sum(
                1 for _, _, t in edges if t == "adjacent"
            )
            logger.info(
                "文档解析完成: %d 条边 (%d 引用 / %d 邻接)",
                len(edges), n_citation, n_adjacent,
            )

        with SimpleTimer("构建 PyG 图"):
            data = build_pyg_graph(
                self.embeddings_original, edges, cfg.device
            ).to(device)
            logger.info(
                "图构建完成: %d 节点, %d 条边",
                data.num_nodes, data.num_edges,
            )

        with SimpleTimer(f"GNN 训练 ({cfg.gnn_model})"):
            try:
                model, best_loss = self._train_gnn(data, d, cfg, device)
            except Exception as exc:
                raise GNNTrainError(f"GNN 训练失败: {exc}") from exc

            logger.info("GNN 训练完成，最佳 Loss: %.6f", best_loss)

        # 推理：生成 GNN 增强 embedding
        model.eval()
        with torch.no_grad():
            gnn_out: torch.Tensor = model(data.x, data.edge_index)
            gnn_emb: torch.Tensor = F.normalize(
                gnn_out * 0.7 + data.x * 0.3, p=2, dim=1
            )

        self.embeddings_gpu = gnn_emb
        self.embeddings_cpu = gnn_emb.detach().cpu().numpy()
        self.gnn_model = model

        self.results["gnn"] = {
            "num_nodes": n,
            "num_edges": data.num_edges,
            "num_citation_edges": n_citation,
            "best_loss": best_loss,
            "model_type": cfg.gnn_model,
            "hidden_dim": cfg.gnn_hidden_dim,
            "num_layers": cfg.gnn_num_layers,
        }

    def _train_gnn(
        self,
        data: "torch_geometric.data.Data",
        in_dim: int,
        cfg: Config,
        device: torch.device,
    ) -> Tuple[GNNEncoder, float]:
        """训练 GNN 模型（自监督重建任务）。

        Args:
            data: PyG Data 对象。
            in_dim: 输入/输出特征维度。
            cfg: 配置对象。
            device: 计算设备。

        Returns:
            (训练好的模型, 最佳 Loss)。
        """
        model: GNNEncoder = GNNEncoder(
            in_dim=in_dim,
            hidden_dim=cfg.gnn_hidden_dim,
            out_dim=in_dim,
            num_layers=cfg.gnn_num_layers,
            dropout=cfg.gnn_dropout,
            model_type=cfg.gnn_model,
        ).to(device)

        optimizer: torch.optim.Adam = torch.optim.Adam(
            model.parameters(), lr=cfg.gnn_lr
        )
        criterion: torch.nn.MSELoss = torch.nn.MSELoss()

        # 验证集划分
        n: int = data.num_nodes
        val_size: int = max(1, int(n * cfg.gnn_val_ratio))
        val_indices: torch.Tensor = torch.randperm(n)[:val_size].to(device)
        train_indices: torch.Tensor = torch.tensor(
            [i for i in range(n) if i not in val_indices],
            dtype=torch.long, device=device,
        )

        model.train()
        best_loss: float = float("inf")
        best_state: Optional[Dict[str, Any]] = None
        patience_counter: int = 0

        for epoch in range(cfg.gnn_epochs):
            model.train()
            optimizer.zero_grad()
            out: torch.Tensor = model(data.x, data.edge_index)
            loss: torch.Tensor = criterion(
                out[train_indices], data.x[train_indices]
            )
            loss.backward()
            optimizer.step()

            # 验证集评估
            model.eval()
            with torch.no_grad():
                val_loss: float = criterion(
                    out[val_indices], data.x[val_indices]
                ).item()

            if val_loss < best_loss:
                best_loss = val_loss
                best_state = {
                    k: v.detach().cpu().clone()
                    for k, v in model.state_dict().items()
                }
                patience_counter = 0
            else:
                patience_counter += 1

            if (epoch + 1) % 50 == 0 or epoch == 0:
                logger.info(
                    "  Epoch %3d/%d | Train Loss: %.6f | Val Loss: %.6f",
                    epoch + 1, cfg.gnn_epochs,
                    loss.item(), val_loss,
                )

            # 早停
            if (
                cfg.gnn_early_stop_patience is not None
                and patience_counter >= cfg.gnn_early_stop_patience
            ):
                logger.info(
                    "  早停触发: Epoch %d (patience=%d)",
                    epoch + 1, cfg.gnn_early_stop_patience,
                )
                break

        if best_state is not None:
            model.load_state_dict(best_state)

        return model, best_loss

    # ------------------------------------------------------------------
    # PHATE 降维
    # ------------------------------------------------------------------

    def run_phate(self, cfg: Config) -> None:
        """运行 PHATE 降维，将 GNN 增强 embedding 映射到 3D 流形空间。

        Args:
            cfg: 配置对象。
        """
        if not _PHATE_AVAILABLE:
            logger.warning("PHATE 不可用，跳过")
            return
        import phate as _phate

        op: _phate.PHATE = _phate.PHATE(
            n_components=cfg.phate_n_components,
            knn=cfg.phate_knn,
            knn_dist=cfg.phate_knn_dist,
            n_pca=cfg.phate_n_pca,
            mds=cfg.phate_mds,
            mds_solver=cfg.phate_mds_solver,
            n_jobs=-1,
            verbose=False,
        )
        self.results["phate"] = op.fit_transform(self.embeddings_cpu)
        self.results["phate_params"] = {
            "n_components": cfg.phate_n_components,
            "knn": cfg.phate_knn,
            "knn_dist": cfg.phate_knn_dist,
            "n_pca": cfg.phate_n_pca,
            "mds": cfg.phate_mds,
            "mds_solver": cfg.phate_mds_solver,
        }

    # ------------------------------------------------------------------
    # TDA 拓扑数据分析
    # ------------------------------------------------------------------

    def run_tda(self) -> None:
        """计算 Vietoris–Rips 持续同调（H0 / H1 / H2）。

        仅在 ``Config.enable_tda`` 为 ``True`` 且 gtda 可用时生效。
        """
        if not _TDA_AVAILABLE:
            return
        from tqdm import tqdm

        from gtda.homology import VietorisRipsPersistence

        n: int = len(self.embeddings_gpu)
        norm: torch.Tensor = F.normalize(self.embeddings_gpu, p=2, dim=1)
        logger.info("计算 %d×%d 距离矩阵 ...", n, n)

        with tqdm(total=3, desc="TDA 步骤") as pbar:
            dist_matrix: np.ndarray = (
                1.0
                - torch.clamp(torch.mm(norm, norm.t()), -1.0, 1.0)
            ).cpu().numpy()
            pbar.update(1)
            pbar.set_description("后处理距离矩阵")
            np.fill_diagonal(dist_matrix, 0.0)
            dist_matrix = np.maximum(dist_matrix, 0.0)
            dist_matrix = (dist_matrix + dist_matrix.T) / 2.0
            pbar.update(1)
            pbar.set_description("VietorisRips 持续同调")
            vr: VietorisRipsPersistence = VietorisRipsPersistence(
                metric="precomputed",
                homology_dimensions=[0, 1, 2],
                collapse_edges=True,
                n_jobs=-1,
            )
            dgms: np.ndarray = vr.fit_transform(
                dist_matrix[np.newaxis, :, :]
            )
            pbar.update(1)

        self.results["tda"] = {
            "diagrams": dgms[0],
            "betti": {0: 0, 1: 0, 2: 0},
        }

    # ------------------------------------------------------------------
    # 离群检测 + 图聚类
    # ------------------------------------------------------------------

    def _knn_graph(
        self, emb: np.ndarray, k: int
    ) -> Tuple[np.ndarray, np.ndarray]:
        """计算 k-NN 距离和索引。

        Args:
            emb: 形状 ``(n, d)`` 的归一化 embedding。
            k: 邻居数量。

        Returns:
            (distances[:, 1:], indices[:, 1:])，自环被排除。
        """
        nn: NearestNeighbors = NearestNeighbors(
            n_neighbors=min(k + 1, len(emb)), metric="cosine"
        )
        nn.fit(emb)
        dists: np.ndarray
        inds: np.ndarray
        dists, inds = nn.kneighbors(emb)
        return dists[:, 1:], inds[:, 1:]

    def run_graph_clustering(self, cfg: Config) -> None:
        """执行离群分析 + Leiden / HDBSCAN 聚类。

        Args:
            cfg: 配置对象。
        """
        emb: np.ndarray = self.embeddings_cpu
        n: int = len(emb)

        if n == 0:
            self.results["clusters"] = np.array([], dtype=int)
            self.results["clusters_raw"] = np.array([], dtype=int)
            self.results["outliers"] = np.array([], dtype=float)
            self.results["cluster_backend"] = "empty"
            return
        if n < 3:
            labels: np.ndarray = np.zeros(n, dtype=int)
            self.results["clusters_raw"] = labels
            self.results["clusters"] = labels
            self.results["outliers"] = np.zeros(n, dtype=float)
            self.results["cluster_backend"] = "short"
            return

        emb_norm: np.ndarray = emb / (
            np.linalg.norm(emb, axis=1, keepdims=True) + 1e-9
        )
        k: int = int(
            np.clip(
                np.sqrt(max(n, 4)),
                cfg.outlier_k_min,
                cfg.outlier_k_max,
            )
        )
        dists, inds = self._knn_graph(emb_norm, k)
        sims: np.ndarray = np.clip(1.0 - dists, 0.0, 1.0)
        outlier_scores: np.ndarray = 1.0 - sims.mean(axis=1)

        labels_raw: Optional[np.ndarray] = None
        backend: str = "unknown"

        if _LEIDEN_AVAILABLE:
            import igraph as _ig
            import leidenalg as _la

            leiden_edges: List[Tuple[int, int]] = []
            weights: List[float] = []
            seen: Set[Tuple[int, int]] = set()
            for i in range(n):
                for j, sim_val in zip(inds[i], sims[i]):
                    a, b = (i, int(j)) if i < int(j) else (int(j), i)
                    if a == b or (a, b) in seen:
                        continue
                    seen.add((a, b))
                    leiden_edges.append((a, b))
                    weights.append(float(sim_val))

            g = _ig.Graph(n=n, edges=leiden_edges, directed=False)
            part = _la.find_partition(
                g,
                _la.RBConfigurationVertexPartition,
                weights=weights,
                resolution_parameter=cfg.leiden_resolution,
            )
            labels_raw = np.array(part.membership, dtype=int)
            backend = "leiden"
        elif _HDBSCAN_AVAILABLE:
            import hdbscan as _hdbscan

            clusterer: _hdbscan.HDBSCAN = _hdbscan.HDBSCAN(
                min_cluster_size=cfg.hdbscan_min_cluster_size,
                min_samples=cfg.hdbscan_min_samples,
            )
            labels_raw = clusterer.fit_predict(emb_norm)
            outlier_scores = clusterer.outlier_scores_
            backend = "hdbscan"
        else:
            labels_raw = np.zeros(n, dtype=int)
            backend = "fallback-single"

        self.results["clusters_raw"] = labels_raw
        self.results["clusters"] = labels_raw
        self.results["outliers"] = outlier_scores
        self.results["cluster_backend"] = backend
        self.results["outlier_k"] = k

# ==================== 模块 6：数据加载 ====================



def load_or_build_chunks(
    segments: List[Dict[str, str]], output_dir: str
) -> List[str]:
    """加载或构建 chunk 列表。

    对法规文档，保留原始段落结构，不过度分割句子。
    如果 ``chunks.json`` 缓存存在则直接加载。

    Args:
        segments: ``split_text_into_segments`` 的输出。
        output_dir: 缓存目录。

    Returns:
        文本列表。
    """
    chunks_cache: str = os.path.join(output_dir, "chunks.json")
    if os.path.exists(chunks_cache):
        with open(chunks_cache, "r", encoding="utf-8") as f:
            chunks: List[str] = json.load(f)
        logger.info("已读取 chunks 缓存: %s", chunks_cache)
        return chunks

    chunks = [s["text"].strip() for s in segments]
    with open(chunks_cache, "w", encoding="utf-8") as f:
        json.dump(chunks, f, ensure_ascii=False)
    logger.info("已保存 chunks 缓存: %s", chunks_cache)
    return chunks


def load_or_embed(
    chunks: List[str], output_dir: str, cfg: Config
) -> torch.Tensor:
    """加载或计算 SBERT embedding。

    优先从 ``embeddings.npy`` 缓存读取；
    否则使用 MLX 或 SentenceTransformer 编码。

    Args:
        chunks: 文本列表。
        output_dir: 缓存目录。
        cfg: 配置对象。

    Returns:
        形状 ``(n, d)`` 的 float32 张量。
    """
    emb_cache: str = os.path.join(output_dir, "embeddings.npy")
    if os.path.exists(emb_cache):
        embeddings: np.ndarray = np.load(emb_cache)
        embeddings_tensor: torch.Tensor = torch.tensor(
            embeddings, dtype=torch.float32
        )
        logger.info("已读取 embeddings 缓存: %s", emb_cache)
        return embeddings_tensor

    # 优先尝试 MLX 加速（Apple Silicon）
    try:
        from mlx_embedding_models.embedding import EmbeddingModel

        try:
            from transformers import PreTrainedTokenizerBase

            if not hasattr(
                PreTrainedTokenizerBase, "batch_encode_plus"
            ):

                def _batch_encode_plus(
                    self: Any,
                    batch_text_or_text_pairs: Any,
                    **kwargs: Any,
                ) -> Any:
                    """补丁：为 PreTrainedTokenizerBase 添加 batch_encode_plus。"""
                    return self.__call__(
                        batch_text_or_text_pairs, **kwargs
                    )

                PreTrainedTokenizerBase.batch_encode_plus = (  # type: ignore[method-assign]
                    _batch_encode_plus
                )
        except Exception:
            pass

        if hasattr(EmbeddingModel, "from_pretrained"):
            mlx_model: EmbeddingModel = EmbeddingModel.from_pretrained(
                cfg.sbert_model_name
            )
        else:
            mlx_model = EmbeddingModel.from_registry("bge-m3")

        embeddings = np.array(mlx_model.encode(chunks))
        logger.info("使用 MLX Embedding Models: %s", cfg.sbert_model_name)
    except Exception as exc:
        logger.warning(
            "MLX 不可用，回退 SentenceTransformer。原因: %s", exc
        )
        model: SentenceTransformer = SentenceTransformer(
            cfg.sbert_model_name, device=cfg.device
        )
        embeddings = model.encode(chunks, convert_to_tensor=True)

    np.save(emb_cache, embeddings)
    logger.info("已保存 embeddings 缓存: %s", emb_cache)
    return torch.tensor(embeddings, dtype=torch.float32)


def compute_coherence(
    embeddings_tensor: torch.Tensor,
) -> Tuple[float, np.ndarray]:
    """计算相邻句子间的余弦相似度（连贯性指标）。

    Args:
        embeddings_tensor: 形状 ``(n, d)`` 的 embedding 张量。

    Returns:
        (平均连贯性, 每个相邻对的相似度数组)。
    """
    norm: torch.Tensor = F.normalize(embeddings_tensor, p=2, dim=1)
    sims: np.ndarray = (
        (norm[:-1] * norm[1:]).sum(dim=1).cpu().numpy()
    )
    return (float(np.mean(sims)), sims)


# ==================== 模块 7：报告生成 ====================


def _build_subplot_specs_and_titles(
    res: Dict[str, Any],
) -> Tuple[List[List[Dict[str, Any]]], List[str]]:
    """根据分析结果动态构建 subplot 布局和标题。

    Args:
        res: 分析结果字典。

    Returns:
        (specs, titles) 用于 ``make_subplots``。
    """
    specs: List[List[Dict[str, Any]]] = [
        [{"type": "scene"}, {"type": "xy"}],
        [{"type": "xy", "colspan": 2}, None],
    ]
    titles: List[str] = [
        "PHATE 逻辑流轨迹 (GNN 增强 + Leiden 聚类)",

        "密度离群分析",
        "语义连贯性实时演变",
    ]

    if "tda" in res:
        specs.append([{"type": "xy"}, {"type": "domain"}])
        titles.append("持续性图 (TDA)")
        titles.append("拓扑逻辑摘要")

    if "phate" in res:
        specs.append([{"type": "domain", "colspan": 2}, None])
        titles.append("PHATE 轨迹特征量化分析")

    specs.append([{"type": "table", "colspan": 2}, None])
    titles.append("Chunk 文本")
    specs.append([{"type": "table", "colspan": 2}, None])
    titles.append("参数清单")


    return specs, titles


def _add_phate_traces(
    fig: go.Figure,
    p_coords: np.ndarray,
    seg_boundaries: Sequence[int],
    seg_topics: Sequence[int],
    phate_point_trace_indices: List[int],
) -> None:
    """添加 PHATE 3D 轨迹相关 trace。

    包括：时间序直线轨迹、分割点连线、各段散点、起止标记。

    Args:
        fig: Plotly Figure 对象。
        p_coords: 形状 ``(n, 3)`` 的 PHATE 坐标。
        seg_boundaries: 分段边界列表。
        seg_topics: 各段对应的主题簇编号。
        phate_point_trace_indices: 用于跟踪散点 trace 索引的列表。
    """
    # 时间序轨迹（直线连接所有点）
    fig.add_trace(
        go.Scatter3d(
            x=p_coords[:, 0],
            y=p_coords[:, 1],
            z=p_coords[:, 2],
            mode="lines",
            line=dict(color="rgba(60,60,60,0.5)", width=1.5),
            name="时间序轨迹",
        ),
        row=1,
        col=1,
    )

    palette: List[str] = [
        "#1f77b4", "#ff7f0e", "#2ca02c", "#d62728", "#9467bd",
        "#8c564b", "#e377c2", "#7f7f7f", "#bcbd22", "#17becf",
    ]

    n: int = len(p_coords)

    # 分割点连线
    if len(seg_boundaries) >= 2:
        b_idx: List[int] = [
            max(0, min(int(b), n - 1)) for b in seg_boundaries
        ]
        fig.add_trace(
            go.Scatter3d(
                x=p_coords[b_idx, 0],
                y=p_coords[b_idx, 1],
                z=p_coords[b_idx, 2],
                mode="lines+markers",
                line=dict(color="rgba(0,0,0,0.6)", width=2.5),
                marker=dict(size=4, color="black", symbol="x"),
                name="分割点连线",
            ),
            row=1,
            col=1,
        )

    # 各段散点
    n_seg: int = max(0, len(seg_boundaries) - 1)
    for s in range(n_seg):
        l: int = int(seg_boundaries[s])
        r: int = int(seg_boundaries[s + 1])
        idx: np.ndarray = np.arange(l, r, dtype=int)
        if len(idx) <= 0:
            continue
        color: str = palette[s % len(palette)]
        topic: int = seg_topics[s] if s < len(seg_topics) else -1
        hover_data: List[List[int]] = [[int(j + 1)] for j in idx]
        fig.add_trace(
            go.Scatter3d(
                x=p_coords[idx, 0],
                y=p_coords[idx, 1],
                z=p_coords[idx, 2],
                mode="markers",
                marker=dict(size=3, color=color, opacity=0.7),
                customdata=hover_data,
                hovertemplate="语义块 %{customdata[0]}<extra></extra>",
                name=f"段 {s + 1} | 主题簇 {int(topic)}",
            ),
            row=1,
            col=1,
        )
        phate_point_trace_indices.append(len(fig.data) - 1)

    # START 标记
    fig.add_trace(
        go.Scatter3d(
            x=[p_coords[0, 0]],
            y=[p_coords[0, 1]],
            z=[p_coords[0, 2]],
            mode="markers+text",
            marker=dict(size=8, color="red", symbol="diamond"),
            text=["START"],
            name="起点",
        ),
        row=1,
        col=1,
    )
    phate_point_trace_indices.append(len(fig.data) - 1)

    # END 标记
    fig.add_trace(
        go.Scatter3d(
            x=[p_coords[-1, 0]],
            y=[p_coords[-1, 1]],
            z=[p_coords[-1, 2]],
            mode="markers+text",
            marker=dict(size=8, color="green", symbol="circle"),
            text=["END"],
            name="终点",
        ),
        row=1,
        col=1,
    )
    phate_point_trace_indices.append(len(fig.data) - 1)


def save_modular_report(
    chunks: List[str],
    analyzer: LogicAnalyzer,
    coh_data: Tuple[float, np.ndarray],
    output_dir: str,
    embeddings_tensor: torch.Tensor,
    cfg: Config,
    doc_title: str = "text_document",
    report_name: str = "phate_logic_report.html",
) -> None:
    """生成交互式 Plotly 报告并保存为 HTML。

    Args:
        chunks: 文本列表。
        analyzer: 已完成全部分析的 LogicAnalyzer 实例。
        coh_data: ``compute_coherence`` 返回的 (均值, 数组)。
        output_dir: 报告保存目录。
        embeddings_tensor: 原始 embedding 张量（用于 R 序列计算）。
        cfg: 配置对象。
        doc_title: 文档标题。
        report_name: 报告 HTML 文件名。
    """
    res: Dict[str, Any] = analyzer.results

    emb: np.ndarray = embeddings_tensor.detach().cpu().numpy()
    emb_norm: np.ndarray = emb / (
        np.linalg.norm(emb, axis=1, keepdims=True) + 1e-9
    )

    # 计算 R 序列（主题相关性）
    phi: float = 0.7
    c_t_list: List[np.ndarray] = [emb_norm[0]]
    for i in range(1, len(emb_norm)):
        c_t_list.append(
            phi * c_t_list[-1] + (1.0 - phi) * emb_norm[i]
        )
    c_t_arr: np.ndarray = np.array(c_t_list)
    c_t_arr = c_t_arr / (
        np.linalg.norm(c_t_arr, axis=1, keepdims=True) + 1e-9
    )
    R: np.ndarray = np.sum(emb_norm * c_t_arr, axis=1)

    specs, titles = _build_subplot_specs_and_titles(res)
    fig: go.Figure = make_subplots(
        rows=len(specs),
        cols=2,
        column_widths=[0.6, 0.4],
        specs=specs,
        subplot_titles=titles,
    )
    phate_point_trace_indices: List[int] = []

    if "phate" in res:
        p_coords: np.ndarray = res["phate"]
        _add_phate_traces(
            fig, p_coords, [0, len(chunks)], [],
            phate_point_trace_indices,
        )


    # 离群度
    fig.add_trace(
        go.Scatter(
            y=res.get("outliers", []),
            mode="markers",
            marker=dict(color="rgba(255,0,0,0.5)", size=3),
            name="离群度",
        ),
        row=1, col=2,
    )

    # 连贯性
    fig.add_trace(
        go.Scatter(
            y=coh_data[1],
            mode="lines",
            fill="tozeroy",
            name="连贯性",
        ),
        row=2, col=1,
    )

    curr_row: int = 3

    # TDA 图
    if "tda" in res:
        dg: np.ndarray = res["tda"]["diagrams"]
        for d in (0, 1, 2):
            pts: np.ndarray = dg[dg[:, 2] == d]
            fig.add_trace(
                go.Scatter(
                    x=pts[:, 0], y=pts[:, 1], mode="markers",
                    name=f"H{d}",
                ),
                row=curr_row, col=1,
            )
        fig.add_trace(
            go.Table(
                header=dict(values=["维度", "含义"]),
                cells=dict(values=[
                    ["H0", "H1", "H2"],
                    ["主题独立性", "逻辑闭环", "高阶空洞/多主题交错"],
                ]),
            ),
            row=curr_row, col=2,
        )
        curr_row += 1

    # PHATE 量化指标
    if "phate" in res:
        cos_sim: np.ndarray = np.sum(
            emb_norm[:-1] * emb_norm[1:], axis=1
        )
        cos_sim = np.clip(cos_sim, -1.0, 1.0)
        d_t: np.ndarray = np.arccos(cos_sim)
        L_vel: float = float(d_t.mean()) if len(d_t) > 0 else 0.0
        mu_v: np.ndarray = emb.mean(axis=0)
        L_vol: float = float(
            np.sqrt(((emb - mu_v) ** 2).sum(axis=1).mean())
        )
        S_coh: float = float(np.mean(cos_sim)) if len(cos_sim) > 0 else 0.0
        z_vel: float = (
            (L_vel - np.mean(d_t)) / (np.std(d_t) + 1e-9)
            if len(d_t) > 0 else 0.0
        )
        vol_series: np.ndarray = np.sqrt(((emb - mu_v) ** 2).sum(axis=1))
        z_vol: float = (
            (L_vol - np.mean(vol_series)) / (np.std(vol_series) + 1e-9)
            if len(vol_series) > 0 else 0.0
        )
        score_load: float = float(S_coh * (0.5 * z_vel + 0.5 * z_vol))
        D_shock: float = (
            float(np.mean(np.abs(R[1:] - R[:-1])))
            if len(R) > 1 else 0.0
        )
        mu_R: float = float(np.mean(R)) if len(R) > 0 else 0.0
        raw_stab: float = mu_R - 0.5 * D_shock
        z_stab: float = (
            (raw_stab - np.mean(R)) / (np.std(R) + 1e-9)
            if len(R) > 0 else 0.0
        )

        fig.add_trace(
            go.Scatter(y=R, mode="lines", name="主题相关性"),
            row=2, col=1,
        )
        fig.add_trace(
            go.Table(
                header=dict(
                    values=["分析指标", "量化数值", "深度解释"],
                    fill_color="darkred",
                    font=dict(color="white"),
                ),
                cells=dict(values=[
                    ["语义推进速率", "语义空间体积", "语义连贯性",
                     "逻辑负载度综合得分", "主题相关性均值", "稳定性得分"],
                    [f"{L_vel:.4f}", f"{L_vol:.4f}", f"{S_coh:.4f}",
                     f"{score_load:.4f}", f"{mu_R:.4f}", f"{z_stab:.4f}"],
                    [r"$L_{vel} = \frac{1}{N-1}\sum \arccos(\cos_t)$",
                     r"$L_{vol} = \sqrt{\frac{1}{N}\sum \lVert v_t-\mu \rVert^2}$",
                     r"$S_{coh} = \frac{1}{N-1}\sum \cos_t$",
                     r"$Score = S_{coh}\cdot(0.5\,Z(L_{vel})+0.5\,Z(L_{vol}))$",
                     r"$\mu_R = \frac{1}{N}\sum R(t)$",
                     r"$S_{stab} = Z(\mu_R - 0.5\,D_{shock})$"],
                ]),
            ),
            row=curr_row, col=1,
        )
        curr_row += 1

    # Chunk 文本表
    chunk_idx: List[int] = list(range(1, len(chunks) + 1))
    fig.add_trace(
        go.Table(
            header=dict(
                values=["Chunk", "文本"],
                fill_color="darkslategray",
                font=dict(color="white"),
            ),
            cells=dict(values=[chunk_idx, chunks], align=["center", "left"]),
        ),
        row=curr_row, col=1,
    )
    curr_row += 1

    # 参数清单表

    phate_params: Dict[str, Any] = res.get("phate_params", {})
    gnn_info: Dict[str, Any] = res.get("gnn", {})
    param_names: List[str] = [
        "Input type", "SBERT_MODEL_NAME",
        "GNN Model", "GNN.hidden_dim", "GNN.num_layers",
        "GNN.dropout", "Graph.citation_edges", "Graph.adj_k",
        "GNN.best_loss",
        "PHATE.n_components", "PHATE.knn", "PHATE.knn_dist",
        "PHATE.n_pca", "PHATE.mds", "PHATE.mds_solver",
        "Cluster backend", "Leiden resolution",
        "HDBSCAN.min_cluster_size", "HDBSCAN.min_samples",
        "ENABLE_TDA",

    ]
    best_loss_val: str = gnn_info.get("best_loss", "n/a")
    if isinstance(best_loss_val, float):
        best_loss_val = f"{best_loss_val:.6f}"
    param_values: List[str] = [
        "text (direct input)", cfg.sbert_model_name,
        str(gnn_info.get("model_type", cfg.gnn_model)),
        str(gnn_info.get("hidden_dim", cfg.gnn_hidden_dim)),
        str(gnn_info.get("num_layers", cfg.gnn_num_layers)),
        str(cfg.gnn_dropout),
        str(gnn_info.get("num_citation_edges", "n/a")),
        str(cfg.graph_adj_k),
        best_loss_val,
        str(phate_params.get("n_components", cfg.phate_n_components)),
        str(phate_params.get("knn", cfg.phate_knn)),
        str(phate_params.get("knn_dist", cfg.phate_knn_dist)),
        str(phate_params.get("n_pca", cfg.phate_n_pca)),
        str(phate_params.get("mds", cfg.phate_mds)),
        str(phate_params.get("mds_solver", cfg.phate_mds_solver)),
        str(res.get("cluster_backend", "unknown")),
        str(cfg.leiden_resolution),
        str(cfg.hdbscan_min_cluster_size),
        str(cfg.hdbscan_min_samples),
        str(cfg.enable_tda),

    ]
    fig.add_trace(
        go.Table(
            header=dict(
                values=["参数", "值"],
                fill_color="darkred",
                font=dict(color="white"),
            ),
            cells=dict(values=[param_names, param_values], align=["left", "left"]),
        ),
        row=curr_row, col=1,
    )

    # 显示/隐藏点的按钮
    n_traces: int = len(fig.data)
    vis_points_on: List[bool] = [True] * n_traces
    vis_points_off: List[bool] = [True] * n_traces
    for idx in phate_point_trace_indices:
        if 0 <= idx < n_traces:
            vis_points_off[idx] = False

    fig.update_layout(
        height=450 * len(specs),
        title=f"📄 [GNN] PHATE 语义拓扑深度逻辑报告 - {doc_title}",
        template="plotly_white",
        updatemenus=[
            dict(
                type="buttons",
                direction="right",
                x=0.01,
                y=1.08,
                showactive=True,
                buttons=[
                    dict(
                        label="显示点",
                        method="update",
                        args=[{"visible": vis_points_on}],
                    ),
                    dict(
                        label="隐藏点",
                        method="update",
                        args=[{"visible": vis_points_off}],
                    ),
                ],
            )
        ],
        scene=dict(
            aspectmode="data",
            xaxis=dict(showbackground=False, gridcolor="#d9e2ef"),
            yaxis=dict(showbackground=False, gridcolor="#d9e2ef"),
            zaxis=dict(showbackground=False, gridcolor="#d9e2ef"),
            camera=dict(eye=dict(x=1.3, y=1.25, z=0.9)),
        ),
    )

    report_path: str = os.path.join(output_dir, report_name)
    fig.write_html(report_path, include_mathjax="cdn")
    logger.info("报告已生成: %s", report_path)


# ==================== 入口 ====================


def build_pipeline(
    text: str,
    cfg: Config,
    texts_original: Optional[List[str]] = None,
) -> None:
    """全自动 GNN 语义分析管线。

    流程:
    1. 文字分段 → 2. 单句 chunks → 3. SBERT 编码
    → 4. **文档图构建 → GNN 传播（替代 EMA）**
    → 5. PHATE 降维 → 6. TDA / 聚类 → 7. 报告生成


    Args:
        text: 要分析的文字内容（拼接后的全文）。
        cfg: 配置对象。
        texts_original: 原始段落列表（保留文档结构，用于图构建）。
    """
    script_dir: str = get_script_dir()
    safe_title: str = (
        cfg.doc_title.replace(" ", "_")
        .replace("/", "_")
        .replace("\\", "_")[:120]
        or "text_document"
    )
    output_dir: str = os.path.join(script_dir, safe_title)
    os.makedirs(output_dir, exist_ok=True)

    with SimpleTimer("文字分段"):
        segments: List[Dict[str, str]] = split_text_into_segments(text)
        logger.info("共分割为 %d 个语义段", len(segments))

    chunks: List[str] = load_or_build_chunks(segments, output_dir)

    with SimpleTimer("SBERT 编码"):
        embeddings_tensor: torch.Tensor = load_or_embed(
            chunks, output_dir, cfg
        )

    analyzer: LogicAnalyzer = LogicAnalyzer(
        embeddings_tensor,
        texts_original if texts_original else chunks,
    )

    with SimpleTimer("文档图构建 + GNN 传播 (替代 EMA)"):
        analyzer.run_gnn_propagation(cfg)

    with SimpleTimer("PHATE 轨迹生成 (GNN 增强后降维)"):
        analyzer.run_phate(cfg)

    if _TDA_AVAILABLE and cfg.enable_tda:
        with SimpleTimer("TDA 拓扑分析"):
            analyzer.run_tda()

    with SimpleTimer("图聚类 (Leiden/HDBSCAN)"):
        analyzer.run_graph_clustering(cfg)

    with SimpleTimer("连贯性计算"):

        coh_data: Tuple[float, np.ndarray] = compute_coherence(
            embeddings_tensor
        )

    save_modular_report(
        chunks, analyzer, coh_data, output_dir,
        embeddings_tensor, cfg, cfg.doc_title,
    )


def parse_args(argv: Optional[List[str]] = None) -> Config:
    """从命令行参数或环境变量解析配置。

    优先级：CLI 参数 > 环境变量 > ``Config`` 默认值。

    Args:
        argv: 命令行参数列表（默认使用 ``sys.argv[1:]``）。

    Returns:
        配置对象。
    """
    parser: argparse.ArgumentParser = argparse.ArgumentParser(
        description="GNN 增强版法规文档语义拓扑分析器",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例:\n"
            "  python3 text_baai_gnn.py --input /path/to/doc.docx\n"
            "  python3 text_baai_gnn.py --input doc.docx --model GAT --epochs 100\n"
        ),
    )
    parser.add_argument(
        "--input", "-i", type=str, default="",
        help="输入文档路径（支持 .txt / .docx / .doc）",
    )
    parser.add_argument(
        "--title", "-t", type=str, default="document",
        help="文档标题（用于缓存目录和报告标题）",
    )
    parser.add_argument(
        "--model", "-m", type=str, default="GCN",
        choices=("GCN", "GAT"),
        help="GNN 模型类型（默认: GCN）",
    )
    parser.add_argument(
        "--epochs", "-e", type=int, default=200,
        help="GNN 最大训练轮数（默认: 200）",
    )
    parser.add_argument(
        "--hidden", type=int, default=256,
        help="GNN 隐藏层维度（默认: 256）",
    )
    parser.add_argument(
        "--layers", type=int, default=2,
        help="GNN 层数（默认: 2）",
    )
    parser.add_argument(
        "--lr", type=float, default=0.01,
        help="GNN 学习率（默认: 0.01）",
    )
    parser.add_argument(
        "--no-citation", action="store_true",
        help="禁用引用边",
    )
    return _build_config_from_args(parser.parse_args(argv))


def _build_config_from_args(args: argparse.Namespace) -> Config:
    """将解析后的 args 转换为 Config 对象。

    Args:
        args: 命令行参数命名空间。

    Returns:
        配置对象。
    """
    cfg: Config = Config()

    if args.input:
        cfg.input_file_path = args.input
    cfg.doc_title = args.title
    cfg.gnn_model = args.model
    cfg.gnn_epochs = args.epochs
    cfg.gnn_hidden_dim = args.hidden
    cfg.gnn_num_layers = args.layers
    cfg.gnn_lr = args.lr
    if args.no_citation:
        cfg.graph_use_citation = False

    return cfg


def main(argv: Optional[List[str]] = None) -> None:
    """主入口：从文件读取文本并运行 GNN 分析。

    支持通过 CLI 参数或 ``Config`` 默认值配置输入路径。

    Args:
        argv: 命令行参数列表（默认使用 ``sys.argv[1:]``）。
    """
    cfg: Config = parse_args(argv)

    # 如果 CLI 未指定输入路径，尝试环境变量和默认值
    input_path: str = cfg.input_file_path
    if not input_path:
        input_path = os.environ.get(
            "GNN_INPUT_FILE", ""
        )

    if not input_path:
        logger.error(
            "未指定输入文件。使用 --input 参数或设置 GNN_INPUT_FILE 环境变量。"
        )
        sys.exit(1)

    full_path: str = input_path
    if not os.path.isabs(input_path):
        full_path = os.path.join(get_script_dir(), input_path)

    try:
        text: str = read_text_file(full_path)
    except DocumentNotFoundError:
        logger.error("输入文件不存在: %s", full_path)
        sys.exit(1)
    except (UnsupportedFormatError, DocumentParseError) as exc:
        logger.error("%s", exc)
        sys.exit(1)

    # 对于 DOCX，保留原始段落结构用于图构建
    texts_original: Optional[List[str]] = None
    if full_path.lower().endswith(".docx"):
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(full_path)
            texts_original = [p.text for p in doc.paragraphs]
            logger.info(
                "保留 %d 个原始 DOCX 段落用于图构建",
                len(texts_original),
            )
        except Exception as exc:
            logger.warning("无法读取 DOCX 段落结构: %s", exc)

    logger.info("已读取输入文件: %s", full_path)
    build_pipeline(text, cfg, texts_original)


if __name__ == "__main__":
    main()
